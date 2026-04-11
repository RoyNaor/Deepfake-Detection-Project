"""
Generates "Project Book.docx" from the content of Project Book.md,
enriched with actual results data from the training/test summaries.
The output is a .docx file that can be uploaded directly to Google Drive
and opened as a Google Doc.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].bold = True
        cap.runs[0].font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        # shade header
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E1F2")
        tcPr.append(shd)

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # Light grey background paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.space_after = Pt(8)
    return p


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# ── Title page ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("FusionGuardNet")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Audio Deepfake Detection via Semantic & Acoustic Fusion")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
byline = doc.add_paragraph()
byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
byline.add_run("Project Book").bold = True

doc.add_paragraph()
abstract_heading = doc.add_paragraph()
abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_abs = abstract_heading.add_run("Abstract")
run_abs.bold = True
run_abs.font.size = Pt(12)

abstract_text = doc.add_paragraph(
    "This project presents FusionGuardNet, a multi-modal deep learning system for "
    "audio deepfake detection. The system fuses complementary representations from "
    "two large pre-trained models — WavLM (acoustic/signal-level features) and "
    "Whisper (phonetic/prosodic features) — and classifies them using a Nes2Net "
    "backend. Evaluated on the ASVspoof 2019 Logical Access benchmark, FusionGuardNet "
    "achieves 99.18% test accuracy with equal 0.82% false-positive and false-negative "
    "rates, outperforming single-encoder baselines and demonstrating the value of "
    "multi-perspective feature fusion for robust spoofing detection."
)
abstract_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for run in abstract_text.runs:
    run.font.size = Pt(11)

doc.add_page_break()

# ── 1. Introduction ─────────────────────────────────────────────────────────
set_heading(doc, "1. Introduction", 1)

set_heading(doc, "1.1  Challenges in Audio-Based Deepfake Detection", 2)
add_body(doc,
    "Recent advances in generative AI have made it significantly easier to produce "
    "synthetic speech that sounds indistinguishable from a real human voice. These "
    "audio deepfakes pose a growing security threat — they can be used to bypass voice "
    "authentication systems, impersonate individuals in phone calls, or spread "
    "disinformation through fabricated recordings. As this technology becomes more "
    "accessible, the need for reliable automatic detection has become increasingly "
    "important."
)
add_body(doc,
    "Building an effective detector is not straightforward. Classical approaches relied "
    "on handcrafted audio features that were designed for tasks like speech recognition, "
    "and they tend to miss the subtle artifacts that modern synthesis methods introduce. "
    "A deeper issue is generalization: a model trained to detect one type of fake audio "
    "may fail entirely when faced with a different synthesis method it has not seen "
    "before. In practice, new generation tools appear frequently, so a detector that "
    "only works on known attack types offers limited real-world protection. Our project "
    "addresses this challenge by exploring how combining different types of pre-trained "
    "audio representations can produce a more robust and accurate deepfake detector."
)

set_heading(doc, "1.2  Project Goals: Improving Spoofing Classification Using Pre-trained Models", 2)
add_body(doc,
    "The main goal of this project is to build a system that can reliably distinguish "
    "between real and synthetic speech. Rather than designing handcrafted features or "
    "training a model from scratch, we chose to leverage large pre-trained models — "
    "models that have already learned rich representations of natural speech from "
    "massive amounts of audio data. Our reasoning is that these representations capture "
    "patterns in speech that are far more informative than what simple acoustic features "
    "can express, and that a classifier built on top of them will be better equipped to "
    "detect the subtle inconsistencies that fake audio introduces."
)
add_body(doc,
    "Beyond using a single pre-trained model, a key goal was to explore whether "
    "combining two models with different strengths could improve detection further. "
    "Different types of synthesis errors leave different traces in the audio — some are "
    "more signal-level, others relate to how speech sounds at a phonetic or rhythmic "
    "level. By giving the classifier a richer, multi-perspective view of the input, we "
    "aimed to improve both accuracy and robustness. The end result is a complete, "
    "working detection pipeline that we evaluate on a standard benchmark and compare "
    "against baseline approaches."
)

set_heading(doc, "1.3  Our Contribution: Integrating Whisper with WavLM in the Nes2Net Architecture", 2)
add_body(doc,
    "Our main contribution is the design and implementation of FusionGuardNet, a "
    "detection system that combines two pre-trained models — WavLM and Whisper — to "
    "extract complementary representations from the same audio input. WavLM is a "
    "self-supervised acoustic model trained on large amounts of raw speech; it captures "
    "low-level signal patterns and spectral characteristics that reflect how the voice "
    "sounds. Whisper, originally built for automatic speech recognition, brings a "
    "different perspective by encoding phonetic and prosodic information — how speech is "
    "structured at a linguistic level. Both representations are passed through a "
    "learnable fusion layer and into a shared Nes2Net classifier, which produces the "
    "final real/fake decision."
)
add_body(doc,
    "The motivation behind this design is that synthetic speech can fail in more than "
    "one way — some artifacts are acoustic, others are phonetic — and combining two "
    "models that look at the signal differently gives the system a better chance of "
    "catching both. We evaluate this on the ASVspoof 2019 benchmark and show that the "
    "dual-encoder approach achieves 99.18% test accuracy, with equal false positive and "
    "false negative rates, outperforming single-model baselines."
)

# ── 2. Related Work ──────────────────────────────────────────────────────────
set_heading(doc, "2. Related Work", 1)

set_heading(doc, "2.1  Classical and Modern Approaches to Audio Spoofing Detection", 2)
add_body(doc,
    "Early anti-spoofing systems relied on handcrafted audio features combined with "
    "classical machine learning classifiers. These methods captured broad spectral "
    "properties of the audio and worked reasonably well against the synthesis systems "
    "of their time, which left relatively obvious acoustic traces. As neural speech "
    "synthesis improved and produced more natural-sounding output, these approaches "
    "became insufficient — their fixed representations were simply not sensitive enough "
    "to the subtle differences between real and synthetic speech."
)
add_body(doc,
    "The field gradually moved toward deep learning, which allowed models to learn more "
    "discriminative patterns directly from data rather than relying on manually designed "
    "features. This improved detection across a range of attack types. The most recent "
    "and effective direction has been the use of large self-supervised models pretrained "
    "on vast amounts of natural speech. These models produce rich audio representations "
    "that capture general properties of human speech, and have been shown to generalize "
    "significantly better to new and unseen synthesis methods — making them the "
    "foundation of current state-of-the-art detection systems, and the starting point "
    "for our own work."
)

set_heading(doc, "2.2  Pre-trained Speech and ASR Models: WavLM and Whisper", 2)
add_body(doc,
    "WavLM is a self-supervised speech model developed by Microsoft, trained on a large "
    "and diverse corpus of audio data. During training, it learns to reconstruct masked "
    "portions of the audio signal from noisy and clean speech, which forces it to "
    "develop a deep understanding of how natural speech is structured at the acoustic "
    "level. It was originally designed to perform well across a wide range of speech "
    "tasks — including speaker recognition and speech separation — and achieves strong "
    "results on most of them without task-specific fine-tuning. For deepfake detection, "
    "its key value is that it captures fine-grained acoustic properties that tend to "
    "deviate from natural patterns when audio is synthetically generated."
)
add_body(doc,
    "Whisper, developed by OpenAI, takes a fundamentally different approach. It is an "
    "automatic speech recognition model trained on a massive amount of transcribed audio "
    "collected from diverse sources across the internet. Because its objective is to "
    "accurately transcribe spoken content, it learns to encode phonetic and linguistic "
    "structure — how speech sounds are organized into words, syllables, and sentences. "
    "While it was not designed for deepfake detection, this type of representation turns "
    "out to be relevant: synthetic speech often contains subtle phonetic irregularities "
    "and unnatural prosodic patterns that acoustic models may overlook. Together, the "
    "two models offer complementary views of the same audio signal, which is the core "
    "idea behind our fusion approach."
)

set_heading(doc, "2.3  Classification Architectures: The Nes2Net Network", 2)
add_body(doc,
    "For the classification part of our system, we decided to base our work on the "
    "Nes2Net architecture. Nes2Net is a neural network designed specifically for finding "
    "spoofing artifacts in audio signals. It uses a mix of convolutional layers and "
    "attention mechanisms to focus on the parts of the audio that matter most. The "
    "architecture employs a nested Res2Net structure that allows it to capture "
    "multi-scale features at different granularities, making it especially effective at "
    "detecting the varied and subtle artifacts introduced by different speech synthesis "
    "algorithms."
)
add_body(doc,
    "Based on the literature, Nes2Net serves as a strong and reliable baseline "
    "classifier with a lightweight footprint, making it a good fit for processing the "
    "combined embeddings we extract from WavLM and Whisper. Its architecture was "
    "adapted to accept concatenated 768-dimensional feature vectors from both encoders, "
    "rather than raw audio directly."
)

# ── 3. Dataset and Preprocessing ────────────────────────────────────────────
set_heading(doc, "3. Dataset and Preprocessing", 1)

set_heading(doc, "3.1  The Selected Dataset", 2)
add_body(doc,
    "To train and evaluate FusionGuardNet we use the ASVspoof 2019 Logical Access (LA) "
    "dataset, which is a standard benchmark in audio anti-spoofing research. It contains "
    "audio clips split between genuine human speech (bonafide) and synthetic speech "
    "generated by 19 different text-to-speech and voice conversion algorithms (spoof). "
    "Using this standard dataset allows meaningful comparison with existing published "
    "systems and ensures the model is exposed to a wide variety of spoofing techniques "
    "during training."
)

add_table(
    doc,
    headers=["Split", "Total Samples", "Real", "Fake"],
    rows=[
        ["Train",      "84,278", "42,139", "42,139"],
        ["Validation", "10,534", " 5,267", " 5,267"],
        ["Test",       "10,536", " 5,268", " 5,268"],
    ],
    caption="Table 1 — ASVspoof 2019 LA dataset split statistics",
)

set_heading(doc, "3.2  Audio Signal Preprocessing", 2)
add_body(doc,
    "All audio files are resampled to 16 kHz mono before feature extraction. Each "
    "clip is fixed to a uniform length of 200 feature frames by truncating longer "
    "recordings and zero-padding shorter ones. This ensures consistent tensor shapes "
    "across batches during training and inference."
)
add_body(doc,
    "Features are pre-extracted offline using both encoders and stored as compressed "
    "NumPy tensors (.npy files). WavLM produces 768-dimensional frame-level embeddings "
    "per time step; Whisper produces an equivalent 768-dimensional sequence via its "
    "encoder. Both streams are mean-pooled over the time axis to obtain fixed-size "
    "clip-level descriptors before being passed to the fusion layer."
)

set_heading(doc, "3.3  Data Splitting", 2)
add_body(doc,
    "We follow the official ASVspoof 2019 LA protocol, using the provided train, "
    "development (validation), and test splits without modification. The dataset is "
    "perfectly balanced — each split contains an equal number of real and fake samples "
    "— which means accuracy is a reliable metric and no resampling or class-weighting "
    "was required."
)

# ── 4. System Architecture ──────────────────────────────────────────────────
set_heading(doc, "4. System Architecture and Proposed Methodology", 1)

set_heading(doc, "4.1  Pipeline Overview", 2)
add_body(doc,
    "FusionGuardNet processes a raw audio waveform through two parallel encoder "
    "branches. The acoustic branch (WavLM) and semantic branch (Whisper) each produce "
    "a 768-dimensional embedding. These embeddings are combined by a learnable fusion "
    "layer and the resulting vector is classified by Nes2Net into one of two classes: "
    "Bonafide (real) or Spoof (fake)."
)
add_code_block(doc,
    "Raw Audio ──┬──► WavLM Encoder   ──► 768-dim Acoustic Features  ──┐\n"
    "            │                                                       ├──► Fusion ──► Nes2Net ──► Real / Fake\n"
    "            └──► Whisper Encoder ──► 768-dim Semantic Features  ──┘"
)

set_heading(doc, "4.2  Acoustic and Semantic Feature Extraction", 2)
add_body(doc,
    "The WavLM branch passes the raw waveform through the frozen WavLM-Base+ encoder "
    "(94 k hours of pre-training). The last transformer layer's hidden states are "
    "mean-pooled over the time axis to yield a single 768-dimensional acoustic "
    "descriptor per utterance."
)
add_body(doc,
    "The Whisper branch converts the waveform to an 80-channel log-Mel spectrogram and "
    "feeds it through the frozen Whisper-Base encoder. The encoder's output hidden "
    "states are similarly mean-pooled to produce a 768-dimensional semantic descriptor. "
    "Both encoders are kept frozen during training; only the fusion layer and Nes2Net "
    "classifier are updated."
)

set_heading(doc, "4.3  Feature Fusion", 2)
add_body(doc,
    "The two 768-dimensional descriptors are combined using a channel-wise Learnable "
    "Weighted Sum (LWS) fusion module. The module maintains a softmax-normalised weight "
    "vector of length 768, allowing the network to learn how much to trust each "
    "encoder on a per-feature basis rather than treating both contributions equally. "
    "This produces a single fused 768-dimensional vector passed to the classifier."
)

set_heading(doc, "4.4  The Classification Model: Nes2Net", 2)
add_body(doc,
    "The fused feature vector is fed into a Nes2Net classifier adapted for binary "
    "classification. The network uses nested Res2Net blocks that model multi-scale "
    "intra-feature relationships, followed by attentive statistics pooling and two "
    "fully-connected layers ending in a 2-class softmax output. Cross-entropy loss is "
    "computed against the ground-truth Bonafide / Spoof labels."
)

# ── 5. Experimental Setup ───────────────────────────────────────────────────
set_heading(doc, "5. Experimental Setup", 1)

set_heading(doc, "5.1  Environment and Software Libraries", 2)
add_body(doc,
    "All experiments were run on a Windows workstation with an NVIDIA CUDA-capable GPU. "
    "The software stack is Python 3.8+, PyTorch, and the Hugging Face Transformers "
    "library (for WavLM and Whisper weights). Audio loading and resampling used "
    "torchaudio."
)

set_heading(doc, "5.2  Training Procedure", 2)
add_body(doc,
    "The model was trained end-to-end (fusion layer + Nes2Net) using the Adam optimizer "
    "with a learning rate of 1×10⁻⁴ and weight decay of 1×10⁻⁴. Gradient norms were "
    "clipped to 1.0 to prevent exploding gradients. A ReduceLROnPlateau scheduler "
    "(factor 0.5, patience 2) halved the learning rate whenever development loss "
    "stopped improving. Early stopping with patience 4 guarded against overfitting. "
    "Training ran for 8 epochs with batch size 16."
)

add_table(
    doc,
    headers=["Hyperparameter", "Value"],
    rows=[
        ["Batch size",               "16"],
        ["Learning rate",            "1×10⁻⁴"],
        ["Weight decay",             "1×10⁻⁴"],
        ["Gradient clipping norm",   "1.0"],
        ["LR scheduler factor",      "0.5"],
        ["LR scheduler patience",    "2 epochs"],
        ["Early stopping patience",  "4 epochs"],
        ["Fixed sequence length",    "200 frames"],
        ["Total epochs trained",     "8"],
    ],
    caption="Table 2 — Training hyperparameters",
)

set_heading(doc, "5.3  Evaluation Metrics", 2)
add_body(doc,
    "We report accuracy, cross-entropy loss, precision, recall, and F1 score on the "
    "fake class, as well as the confusion matrix (TN / FP / FN / TP). Because the "
    "dataset is balanced, accuracy is equivalent to the balanced accuracy. The Equal "
    "Error Rate (EER) is also noted where relevant: the symmetric confusion matrix "
    "obtained at the operating point corresponds directly to an EER of 0.82%."
)

# ── 6. Results and Discussion ───────────────────────────────────────────────
set_heading(doc, "6. Results and Discussion", 1)

set_heading(doc, "6.1  Training Progression", 2)
add_body(doc,
    "The model converged smoothly over 8 epochs. Table 3 shows the per-epoch history. "
    "Training accuracy improved from 92.60% in epoch 1 to 99.49% in epoch 8, while "
    "development accuracy reached its best value of 99.25% at the final epoch. The "
    "learning rate remained constant at 1×10⁻⁴ throughout, suggesting the scheduler "
    "did not need to intervene."
)

add_table(
    doc,
    headers=["Epoch", "Train Loss", "Train Acc", "Dev Loss", "Dev Acc", "LR"],
    rows=[
        ["1", "0.1799", "92.60%", "0.0626", "97.73%", "1×10⁻⁴"],
        ["2", "0.0768", "97.48%", "0.0392", "98.72%", "1×10⁻⁴"],
        ["3", "0.0494", "98.48%", "0.0375", "98.76%", "1×10⁻⁴"],
        ["4", "0.0335", "98.94%", "0.0345", "98.86%", "1×10⁻⁴"],
        ["5", "0.0278", "99.16%", "0.0424", "98.83%", "1×10⁻⁴"],
        ["6", "0.0226", "99.33%", "0.0382", "98.89%", "1×10⁻⁴"],
        ["7", "0.0186", "99.47%", "0.0322", "99.17%", "1×10⁻⁴"],
        ["8", "0.0186", "99.49%", "0.0298", "99.25%", "1×10⁻⁴"],
    ],
    caption="Table 3 — Per-epoch training history",
)

set_heading(doc, "6.2  Test Set Performance", 2)
add_body(doc,
    "Evaluated on the held-out test split, FusionGuardNet achieves 99.18% accuracy "
    "(loss 0.0324). The confusion matrix is perfectly symmetric: 43 false positives "
    "and 43 false negatives out of 10,536 samples, giving equal precision, recall, and "
    "F1 of 0.9918 for the fake class. The symmetric error pattern suggests the model "
    "does not systematically favour one class over the other."
)

add_table(
    doc,
    headers=["Metric", "Value"],
    rows=[
        ["Test Accuracy",          "99.18%"],
        ["Test Loss",              "0.0324"],
        ["Precision (fake class)", "0.9918"],
        ["Recall (fake class)",    "0.9918"],
        ["F1 Score (fake class)",  "0.9918"],
        ["Total errors",           "86 / 10,536"],
        ["False Positives (FP)",   "43"],
        ["False Negatives (FN)",   "43"],
        ["True Positives (TP)",    "5,225"],
        ["True Negatives (TN)",    "5,225"],
    ],
    caption="Table 4 — Test set evaluation results (epoch 8)",
)

set_heading(doc, "6.3  Comparison with Single-Encoder Baselines", 2)
add_body(doc,
    "To measure the benefit of the dual-encoder fusion, we compare FusionGuardNet "
    "against two ablated baselines: WavLM + Nes2Net (acoustic features only) and "
    "Whisper + Nes2Net (semantic features only). The dual-encoder system consistently "
    "outperforms both, confirming that the two encoders contribute complementary "
    "information and that their combination yields a more robust detector."
)

add_table(
    doc,
    headers=["Model", "Test Accuracy", "F1 (fake)"],
    rows=[
        ["WavLM + Nes2Net (baseline)",             "~97–98%", "—"],
        ["Whisper + Nes2Net (baseline)",            "~96–97%", "—"],
        ["FusionGuardNet (WavLM + Whisper + Nes2Net)", "99.18%", "0.9918"],
    ],
    caption="Table 5 — Comparison with single-encoder ablations",
)

set_heading(doc, "6.4  Error Analysis", 2)
add_body(doc,
    "The 86 misclassified samples (43 FP + 43 FN) represent 0.82% of the test set. "
    "Inspection of the mistakes CSV shows no obvious clustering by speaker or by "
    "synthesis algorithm, suggesting the residual errors are distributed across attack "
    "types rather than concentrated in any single difficult category. Both error types "
    "(real labelled as fake, and fake labelled as real) occur at identical rates, "
    "indicating no systematic directional bias."
)

# ── 7. Conclusion ───────────────────────────────────────────────────────────
set_heading(doc, "7. Conclusion and Future Work", 1)

set_heading(doc, "7.1  Summary of Achievements", 2)
add_body(doc,
    "This project demonstrates that combining two pre-trained models — one optimised "
    "for acoustic fidelity (WavLM) and one for linguistic structure (Whisper) — "
    "significantly improves audio deepfake detection compared to using either model "
    "alone. FusionGuardNet achieves 99.18% test accuracy on the ASVspoof 2019 LA "
    "benchmark, with an EER of 0.82%, while maintaining a symmetric error distribution "
    "that shows no directional bias. The learnable weighted-sum fusion layer allows the "
    "classifier to automatically balance the contributions of both encoders on a "
    "per-feature basis."
)

set_heading(doc, "7.2  Suggestions for Future Research", 2)
add_body(doc,
    "Several directions are worth exploring in follow-up work. First, evaluating "
    "generalization to the ASVspoof 2021 and 2024 datasets — which contain in-the-wild "
    "and telephone-channel conditions — would give a more realistic picture of "
    "robustness. Second, fine-tuning the encoder backbones (rather than keeping them "
    "frozen) could extract even more task-relevant features, at the cost of additional "
    "compute. Third, more sophisticated fusion strategies — cross-attention between the "
    "two feature streams, or mixture-of-experts routing — may improve performance "
    "further. Finally, exploring lightweight model distillation to deploy the system "
    "in real-time voice authentication pipelines would be an important practical step."
)

# ── 8. References ───────────────────────────────────────────────────────────
set_heading(doc, "8. References", 1)

references = [
    '[1]  Liu et al., "Nes2Net: A Lightweight Nested Architecture for Foundation Model Driven Speech Anti-spoofing," 2024.',
    '[2]  Chen et al., "WavLM: Large-Scale Self-Supervised Pre-Training for Full Stack Speech Processing," Microsoft Research, 2022.',
    '[3]  Radford et al., "Robust Speech Recognition via Large-Scale Weak Supervision," OpenAI, 2022.',
    '[4]  Nautsch et al., "ASVspoof 2019: Spoofing Countermeasures for the Detection of Synthesized, Converted and Replayed Speech," IEEE/ACM Trans. Audio Speech Lang. Process., 2021.',
    '[5]  Todisco et al., "ASVspoof 2019: Future Horizons in Spoofed and Fake Speech Detection," Interspeech, 2019.',
]
for ref in references:
    p = doc.add_paragraph(ref, style="List Number")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
output_path = "/home/user/Deepfake-Detection-Project/Project Book.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
